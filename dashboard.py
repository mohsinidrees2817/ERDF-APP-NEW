# dashboard.py
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import re

section_titles = [
    "Full Document Preview",
    "Project Summary",
    "Challenges and Needs",
    "Target Group",
    "Organisation Structure",
    "Risk Analysis",
    "Communication Plan",
    "Internal Policies",
]


def clean_content_for_display(content):
    """Clean content for display purposes only - NOT for DOCX export"""
    if not content:
        return ""
    content = re.sub(r"\*\*Your input:\*\*.*?\n\n", "", content, flags=re.DOTALL)
    content = re.sub(r"\*\*AI-generated draft for .*?:\*\*\n\n", "", content)
    return content.strip()


def get_raw_content(section_name, section_index):
    """Get raw content without any cleaning for DOCX export"""
    # First try to get edited content
    content = st.session_state.edited_sections.get(section_name, "")

    # If no edited content, get generated content
    if not content:
        content = st.session_state.get(f"step_{section_index}_generated", "")

    # If still no content, return placeholder
    if not content:
        return "*No content provided for this section*"

    return content


def format_section_content_for_display(content):
    """Format content for display in Streamlit"""
    cleaned = clean_content_for_display(content)
    if not cleaned:
        return "*No content provided for this section*"
    return cleaned


# def add_paragraph_with_formatting(doc, text):
#     """Add paragraph with proper formatting, handling bold, italic, etc."""
#     if not text.strip():
#         return

#     # Handle different paragraph types
#     if (
#         text.strip().startswith("**")
#         and text.strip().endswith("**")
#         and text.count("**") == 2
#     ):
#         # Full bold paragraph
#         p = doc.add_paragraph()
#         run = p.add_run(text.strip()[2:-2])
#         run.bold = True
#         return

#     # Handle numbered lists - FIXED VERSION
#     if any(text.strip().startswith(f"{i}.") for i in range(1, 20)):
#         # Extract the text after the number and dot
#         match = re.match(r"^\d+\.\s*(.*)", text.strip())
#         if match:
#             list_text = match.group(1)
#             p = doc.add_paragraph(list_text, style="List Number")
#         else:
#             # Fallback if regex fails
#             p = doc.add_paragraph(text.strip(), style="List Number")
#         return

#     # Handle bullet points
#     if text.strip().startswith(("*", "-")) and not text.strip().startswith("**"):
#         # Remove the bullet and add as list item
#         clean_text = text.strip()[1:].strip()
#         p = doc.add_paragraph(clean_text, style="List Bullet")
#         return

#     # Regular paragraph with inline formatting
#     paragraph = doc.add_paragraph()

#     # Split text by markdown formatting
#     parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)

#     for part in parts:
#         if not part:
#             continue
#         if part.startswith("**") and part.endswith("**") and len(part) > 4:
#             # Bold text
#             run = paragraph.add_run(part[2:-2])
#             run.bold = True
#         elif (
#             part.startswith("*")
#             and part.endswith("*")
#             and len(part) > 2
#             and not part.startswith("**")
#         ):
#             # Italic text
#             run = paragraph.add_run(part[1:-1])
#             run.italic = True
#         elif part.startswith("`") and part.endswith("`"):
#             # Code text
#             run = paragraph.add_run(part[1:-1])
#             run.font.name = "Courier New"
#         else:
#             # Regular text
#             paragraph.add_run(part)


def add_markdown_table_to_doc(doc, markdown_text):
    """Convert markdown table to Word table with proper formatting"""
    lines = [line.strip() for line in markdown_text.strip().splitlines()]
    table_lines = [line for line in lines if "|" in line and line.strip() != ""]

    if len(table_lines) < 2:
        add_paragraph_with_formatting(doc, markdown_text)
        return

    # Find header and separator
    header_line = table_lines[0]
    separator_found = False
    data_start_idx = 1

    # Look for separator line (contains dashes)
    for i, line in enumerate(table_lines[1:], 1):
        if "-" in line and all(c in "|-: " for c in line):
            separator_found = True
            data_start_idx = i + 1
            break

    if not separator_found:
        # No separator found, treat first line as header anyway
        data_start_idx = 1

    # Extract headers
    headers = [cell.strip() for cell in header_line.strip("|").split("|")]
    headers = [h for h in headers if h]  # Remove empty headers

    if not headers:
        add_paragraph_with_formatting(doc, markdown_text)
        return

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        if i < len(hdr_cells):
            hdr_cells[i].text = header
            # Make header bold
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    # Add data rows
    for line in table_lines[data_start_idx:]:
        cols = [cell.strip() for cell in line.strip("|").split("|")]
        cols = [
            c for c in cols if c or len(cols) == len(headers)
        ]  # Keep empty cells if row length matches

        if len(cols) != len(headers):
            continue  # skip malformed rows

        row_cells = table.add_row().cells
        for i, cell_content in enumerate(cols):
            if i < len(row_cells):
                row_cells[i].text = cell_content


def process_content_for_docx(doc, content):
    """Process content and add to document with proper formatting"""
    if not content or content.strip() == "*No content provided for this section*":
        doc.add_paragraph("No content provided for this section")
        return

    # Split content into blocks (separated by double newlines)
    blocks = content.split("\n\n")

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # Check if this block is a table
        lines = block.split("\n")
        table_lines = [line for line in lines if "|" in line]

        if len(table_lines) >= 2:
            # This is likely a table block
            add_markdown_table_to_doc(doc, block)
        else:
            # Regular text block - process line by line for better control
            lines = [line.strip() for line in block.split("\n") if line.strip()]

            i = 0
            while i < len(lines):
                line = lines[i]

                # Check for headings
                if line.startswith("###"):
                    heading_text = line[3:].strip()
                    doc.add_heading(heading_text, level=3)
                    i += 1
                elif line.startswith("##"):
                    heading_text = line[2:].strip()
                    doc.add_heading(heading_text, level=2)
                    i += 1
                elif line.startswith("#"):
                    heading_text = line[1:].strip()
                    doc.add_heading(heading_text, level=1)
                    i += 1

                # Check for section headings (like "2.1 - Project Concept")
                elif re.match(r"^\d+\.\d+\s*-\s*.+", line):
                    # This is a subsection heading
                    heading_text = re.sub(r"^\d+\.\d+\s*-\s*", "", line)
                    doc.add_heading(heading_text, level=2)
                    i += 1

                # Check for main section headings (like "2 - Project Idea")
                elif re.match(r"^\d+\s*-\s*.+", line):
                    # This is a main section heading
                    heading_text = re.sub(r"^\d+\s*-\s*", "", line)
                    doc.add_heading(heading_text, level=1)
                    i += 1

                # Check for numbered lists (consecutive numbered items)
                elif re.match(r"^\d+\.\s+", line):
                    # Look ahead to see if this is part of a numbered list
                    j = i
                    list_items = []
                    while j < len(lines) and re.match(r"^\d+\.\s+", lines[j]):
                        # Extract text after the number
                        text_content = re.sub(r"^\d+\.\s+", "", lines[j])
                        list_items.append(text_content)
                        j += 1

                    # Add the numbered list items
                    for item in list_items:
                        p = doc.add_paragraph(item, style="List Number")

                    i = j  # Skip the processed lines

                # Check for bullet points
                elif line.startswith(("*", "-")) and not line.startswith("**"):
                    # Look ahead for consecutive bullet points
                    j = i
                    bullet_items = []
                    while (
                        j < len(lines)
                        and lines[j].startswith(("*", "-"))
                        and not lines[j].startswith("**")
                    ):
                        clean_text = lines[j][1:].strip()
                        bullet_items.append(clean_text)
                        j += 1

                    # Add the bullet list items
                    for item in bullet_items:
                        p = doc.add_paragraph(item, style="List Bullet")

                    i = j  # Skip the processed lines

                # Check for bold paragraphs
                elif (
                    line.startswith("**")
                    and line.endswith("**")
                    and line.count("**") == 2
                ):
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                    i += 1

                # Regular paragraph
                else:
                    add_paragraph_with_formatting(doc, line)
                    i += 1


def add_paragraph_with_formatting(doc, text):
    """Add paragraph with proper formatting, handling bold, italic, etc."""
    if not text.strip():
        return

    # Regular paragraph with inline formatting
    paragraph = doc.add_paragraph()

    # Split text by markdown formatting
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)

    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif (
            part.startswith("*")
            and part.endswith("*")
            and len(part) > 2
            and not part.startswith("**")
        ):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            # Regular text
            paragraph.add_run(part)


def dashboard_ui():
    user = st.session_state.get("user", "guest@example.com")
    st.markdown(
        f"""
        <div style="background-color: #f5f5f5; padding: 1rem; border-radius: 8px; margin-bottom: 1.5rem;">
            <h2>📋 ERDF Application Dashboard</h2>
            <p>👤 Logged in as: <strong>{user}</strong></p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.sidebar.title("Navigation")
    selected_section = st.sidebar.radio("Sections", section_titles)

    if selected_section == "Full Document Preview":
        st.subheader("📄 Complete Application Document")
        st.markdown("---")

        with st.container():
            for i, section in enumerate(section_titles[1:], 1):
                st.markdown(f"## {i}. {section}")
                content = st.session_state.edited_sections.get(
                    section, st.session_state.get(f"step_{i-1}_generated", "")
                )
                st.markdown(format_section_content_for_display(content))
                st.markdown("---")

        with st.expander("✏️ Edit Sections"):
            selected_edit = st.selectbox("Select section to edit:", section_titles[1:])
            section_index = section_titles.index(selected_edit) - 1
            content = st.session_state.edited_sections.get(
                selected_edit,
                st.session_state.get(f"step_{section_index}_generated", ""),
            )
            edited = st.text_area(
                f"Edit {selected_edit}",
                value=content,
                height=200,
                key=f"full_preview_edit_{section_index}",
            )
            if st.button(f"💾 Save changes to {selected_edit}"):
                st.session_state.edited_sections[selected_edit] = edited
                st.success(f"Changes to {selected_edit} saved!")
                st.rerun()

        st.markdown("---")
        st.subheader("Export Options")

        if st.button("⬇️ Download as DOCX"):
            # Create document
            doc = Document()

            # Add title
            title = doc.add_heading("ERDF Application", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add a line break
            doc.add_paragraph()

            # Add each section
            for i, section_name in enumerate(section_titles[1:]):
                # Add section heading
                doc.add_heading(f"{i+1}. {section_name}", level=1)

                # Get raw content (without cleaning)
                content = get_raw_content(section_name, i)

                # Process and add content
                process_content_for_docx(doc, content)

                # Add space between sections
                doc.add_paragraph()

            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Download DOCX",
                buffer.getvalue(),
                file_name="ERDF_Application.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        return

    # Single section view
    st.subheader(f"✏️ {selected_section}")
    section_index = section_titles.index(selected_section) - 1
    content = st.session_state.edited_sections.get(
        selected_section,
        st.session_state.get(
            f"step_{section_index}_generated", "No content available yet."
        ),
    )
    edited_content = st.text_area(
        "Edit this section:", value=content, height=400, key=f"edit_{selected_section}"
    )
    if st.button("💾 Save Changes"):
        st.session_state.edited_sections[selected_section] = edited_content
        st.success("Changes saved to your application!")

    st.divider()
    st.subheader("Your Original Input")
    user_input = st.session_state.get(
        f"step_{section_index}_input", "No input provided."
    )
    st.info(user_input)


dashboard_ui()
