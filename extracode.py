# dashboard.py
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO
import re

section_titles = [
    "Full Document Preview",
    "Project Summary",
    "Challenges and Needs",
    "Target Group",
    "Organisation Structure",
    "Risk Analysis",
    "Communication Plan",
    "Internal Policies",
]


def clean_content_for_display(content):
    """Clean content for display purposes only - NOT for DOCX export"""
    if not content:
        return ""
    content = re.sub(r"\*\*Your input:\*\*.*?\n\n", "", content, flags=re.DOTALL)
    content = re.sub(r"\*\*AI-generated draft for .*?:\*\*\n\n", "", content)
    return content.strip()


def get_raw_content(section_name, section_index):
    """Get raw content without any cleaning for DOCX export"""
    # First try to get edited content
    content = st.session_state.edited_sections.get(section_name, "")

    # If no edited content, get generated content
    if not content:
        content = st.session_state.get(f"step_{section_index}_generated", "")

    # If still no content, return placeholder
    if not content:
        return "*No content provided for this section*"

    return content


def format_section_content_for_display(content):
    """Format content for display in Streamlit"""
    cleaned = clean_content_for_display(content)
    if not cleaned:
        return "*No content provided for this section*"
    return cleaned


def add_paragraph_with_formatting(doc, text):
    """Add paragraph with proper formatting, handling bold, italic, etc."""
    paragraph = doc.add_paragraph()

    # Split text by markdown formatting
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)

    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            # Code text
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
        else:
            # Regular text
            paragraph.add_run(part)


def add_markdown_table_to_doc(doc, markdown_text):
    """Convert markdown table to Word table with proper formatting"""
    lines = [line.strip() for line in markdown_text.strip().splitlines() if "|" in line]

    if len(lines) < 2:
        add_paragraph_with_formatting(doc, markdown_text)
        return

    # Extract headers
    headers = [cell.strip() for cell in lines[0].strip("|").split("|")]

    # Create table
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    # Add headers
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Make header bold
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add data rows (skip separator line)
    for line in lines[2:]:
        cols = [cell.strip() for cell in line.strip("|").split("|")]
        if len(cols) != len(headers):
            continue  # skip malformed rows

        row_cells = table.add_row().cells
        for i, cell_content in enumerate(cols):
            row_cells[i].text = cell_content


def process_content_for_docx(doc, content):
    """Process content and add to document with proper formatting"""
    if not content or content.strip() == "*No content provided for this section*":
        doc.add_paragraph("No content provided for this section")
        return

    # Check if content contains a table
    if "|" in content and "-" in content:
        # Split content into parts (text before table, table, text after table)
        parts = content.split("\n")
        table_start = -1
        table_end = -1

        for i, line in enumerate(parts):
            if "|" in line and table_start == -1:
                table_start = i
            elif "|" in line:
                table_end = i

        if table_start != -1:
            # Add content before table
            if table_start > 0:
                before_table = "\n".join(parts[:table_start]).strip()
                if before_table:
                    for paragraph in before_table.split("\n\n"):
                        if paragraph.strip():
                            add_paragraph_with_formatting(doc, paragraph.strip())

            # Add table
            table_content = "\n".join(parts[table_start : table_end + 1])
            add_markdown_table_to_doc(doc, table_content)

            # Add content after table
            if table_end < len(parts) - 1:
                after_table = "\n".join(parts[table_end + 1 :]).strip()
                if after_table:
                    for paragraph in after_table.split("\n\n"):
                        if paragraph.strip():
                            add_paragraph_with_formatting(doc, paragraph.strip())
    else:
        # Regular text content
        paragraphs = content.split("\n\n")
        for paragraph in paragraphs:
            if paragraph.strip():
                add_paragraph_with_formatting(doc, paragraph.strip())


def dashboard_ui():
    user = st.session_state.get("user", "guest@example.com")
    st.markdown(
        f"""
        <div style="background-color: #f5f5f5; padding: 1rem; border-radius: 8px; margin-bottom: 1.5rem;">
            <h2>📋 ERDF Application Dashboard</h2>
            <p>👤 Logged in as: <strong>{user}</strong></p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    st.sidebar.title("Navigation")
    selected_section = st.sidebar.radio("Sections", section_titles)

    if selected_section == "Full Document Preview":
        st.subheader("📄 Complete Application Document")
        st.markdown("---")

        with st.container():
            for i, section in enumerate(section_titles[1:], 1):
                st.markdown(f"## {i}. {section}")
                content = st.session_state.edited_sections.get(
                    section, st.session_state.get(f"step_{i-1}_generated", "")
                )
                st.markdown(format_section_content_for_display(content))
                st.markdown("---")

        with st.expander("✏️ Edit Sections"):
            selected_edit = st.selectbox("Select section to edit:", section_titles[1:])
            section_index = section_titles.index(selected_edit) - 1
            content = st.session_state.edited_sections.get(
                selected_edit,
                st.session_state.get(f"step_{section_index}_generated", ""),
            )
            edited = st.text_area(
                f"Edit {selected_edit}",
                value=content,
                height=200,
                key=f"full_preview_edit_{section_index}",
            )
            if st.button(f"💾 Save changes to {selected_edit}"):
                st.session_state.edited_sections[selected_edit] = edited
                st.success(f"Changes to {selected_edit} saved!")
                st.rerun()

        st.markdown("---")
        st.subheader("Export Options")

        if st.button("⬇️ Download as DOCX"):
            # Create document
            doc = Document()

            # Add title
            title = doc.add_heading("ERDF Application", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add a line break
            doc.add_paragraph()

            # Add each section
            for i, section_name in enumerate(section_titles[1:]):
                # Add section heading
                doc.add_heading(f"{i+1}. {section_name}", level=1)

                # Get raw content (without cleaning)
                content = get_raw_content(section_name, i)

                # Process and add content
                process_content_for_docx(doc, content)

                # Add space between sections
                doc.add_paragraph()

            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.download_button(
                "📥 Download DOCX",
                buffer.getvalue(),
                file_name="ERDF_Application.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        return

    # Single section view
    st.subheader(f"✏️ {selected_section}")
    section_index = section_titles.index(selected_section) - 1
    content = st.session_state.edited_sections.get(
        selected_section,
        st.session_state.get(
            f"step_{section_index}_generated", "No content available yet."
        ),
    )
    edited_content = st.text_area(
        "Edit this section:", value=content, height=400, key=f"edit_{selected_section}"
    )
    if st.button("💾 Save Changes"):
        st.session_state.edited_sections[selected_section] = edited_content
        st.success("Changes saved to your application!")

    st.divider()
    st.subheader("Your Original Input")
    user_input = st.session_state.get(
        f"step_{section_index}_input", "No input provided."
    )
    st.info(user_input)


dashboard_ui()
